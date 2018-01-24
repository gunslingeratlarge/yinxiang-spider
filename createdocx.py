#!/usr/bin/env python
#! -*- coding:utf-8 -*-

from docx import Document
from docx.oxml.ns import qn
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.text.run import Run

#这个星期是一个比较充实的星期，我初步了解了xxx的使用，熟悉了xxx，并对xxx有了更为深入的认识。同时还进行了xxx的学习，并结合xxx"
#来不断充实和完善自己。除此之外，在xxx方面也进行了一些关注，包括xxx,xxx,xxx以及xxx，尤其是其中的xxx等内容。总之，这个星期是一个富有收获的星期，学有所思，学有所得


class MyDoc:
    def __init__(self):
        self.doc = Document()
        self.doc.styles['Normal'].font.name = u'宋体'
        self.doc.styles['Normal']._element.rPr.rFonts.set(qn('w:eastAsia'), u'宋体')
        self.doc.styles['Heading 2'].font.name = u'微软雅黑'
        self.doc.styles['Heading 2']._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
        self.runlist = []

    def add_title(self, title):
       h1 = self.doc.add_heading()
       h1.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
       run = h1.add_run(title.decode("utf-8"))
       self.runlist.append(run)

    def add_heading2(self,heading):
        h2 = self.doc.add_heading(level=2)
        run = h2.add_run(heading.decode("utf-8"))
        self.runlist.append(run)

    def add_para(self,para):
        para = self.doc.add_paragraph(para.decode("utf-8"))

    def set_font(self):
        for run in self.runlist:
            run.font.name = u'微软雅黑'
            run.font.bold = 0
            run._element.rPr.rFonts.set(qn('w:eastAsia'), u'微软雅黑')
    def save(self,name):
        self.set_font()
        self.doc.save(name.decode("utf-8"))

if __name__ == "__main__":
    mydoc = MyDoc()
    mydoc.add_title("周报告")
    mydoc.add_heading2("一、概述")
    mydoc.add_para("这个星期是一个比较充实的星期，我初步了解了xxx的使用，熟悉了xxx，并对xxx有了更为深入的认识。同时还进行了xxx的学习，并结合xxx")
    mydoc.add_heading2("二、具体内容")
    mydoc.save("周报告.docx")

